#import libaries

import pandas as pd
import os
import numpy as np
import matplotlib
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import statistics 
import seaborn as sns
from bokeh.plotting import figure, show, output_file, output_notebook, save
from bokeh.models import HoverTool, ColumnDataSource
from PIL import Image,ImageDraw
from matplotlib.colors import Normalize
from matplotlib.colors import ListedColormap
import gc  
import warnings
warnings.simplefilter(action='ignore', category=FutureWarning)

#setup figure figure format

font = {'family' : 'Arial',
        'weight' : 'bold',
        'size'   : 22}
Size = 30

matplotlib.rc('font', **font)
matplotlib.rc('font', size=Size)         
matplotlib.rc('axes', titlesize=Size)     
matplotlib.rc('axes', labelsize=Size)   
matplotlib.rc('xtick', labelsize=Size)    
matplotlib.rc('ytick', labelsize=Size)    
matplotlib.rc('legend', fontsize=Size)    
matplotlib.rc('figure', titlesize=Size)  

'''patchclampdata object - the aim of this object is to take a location of a CSV file contain patch clamp data and add functionallity for thresholding and plotting
    this is very much a first run at somthing like this so there is lots of code cleaning and optimisation that can be made
    though we hope it can act as a starting point for people implimenting this kind of method 
    of analysis to patch clampe data.
    
    The function was writen to deal with fixed voltage postive current experiments though has been used to analyse all data in the paper using some work arounds
    
    This is very much a proof of concept and though it can generate the figures there is room for improvment in code quality, commenting and implimnetation'''

class PatchClampData:
    instancelist = []
    experiment_summary_dataframe = pd.DataFrame()
    count = 0
    def __init__(self, file_path,exp_cond = None ,molecule_id=None, time_col_to_sort='Time[s](Relative to experiment start)', molecule_col='Molecule_ID', exp_id_col='Experiment_ID', exp_cond_col='Exp_Conditions', electrical_current_col='I-mon[A]', time_col='Time[s](Relative to experiment start)', row_data_is_extracted=0):
        try:
            PatchClampData.instancelist.append(self)
            PatchClampData.count +=1
            self.patch_clamp_data = pd.read_csv(file_path)
            self.patch_clamp_data = self.patch_clamp_data.sort_values(by=time_col_to_sort)
            
            # Validate that the specified columns exist
            for col in [molecule_col, exp_id_col, exp_cond_col, electrical_current_col]:
                if col not in self.patch_clamp_data.columns:
                    raise ValueError(f"Column '{col}' not found in the dataset.")
            #deal with nan columns in electrical_current_col
            initial_rows = len(self.patch_clamp_data)
            self.patch_clamp_data = self.patch_clamp_data.dropna(subset=[electrical_current_col])
            rows_dropped = initial_rows - len(self.patch_clamp_data)
            

            # Print a message if rows are dropped due to NaN values
            if rows_dropped > 0:
                print(f"{rows_dropped} row(s) dropped due to NaN values in column '{electrical_current_col}'.")
                
            
            # Initialize attributes
            if molecule_id is None:
                self.molecule_ID = self.patch_clamp_data[molecule_col][row_data_is_extracted]
            else:
                self.molecule_ID = molecule_id
            if exp_cond is None:
                self.exp_cond = self.patch_clamp_data[exp_cond_col][row_data_is_extracted].replace('_',' ')
            else:
                self.exp_cond = exp_cond
            
            self.exp_ID = self.patch_clamp_data[exp_id_col][row_data_is_extracted]
            self.electrical_current_col = electrical_current_col
            self.time_col = time_col
            self.length_of_exp = self.patch_clamp_data[self.time_col].max()

            # naming functionality specfic to changing interanl names to paper naming
            self.plot_name = r"$\mathbf{" + str(self.molecule_ID) + "}$" + f" ({self.exp_cond}), experiment id {self.exp_ID}"
            
            if self.molecule_ID == '3 and 5':
                self.plot_name = r"$\mathbf{3}$ and $\mathbf{5}$" + f" ({self.exp_cond}), experiment id {self.exp_ID}"
            if self.molecule_ID == '7 and 9':
                self.plot_name = r"$\mathbf{7}$ and $\mathbf{9}$" + f" ({self.exp_cond}), experiment id {self.exp_ID}"
            
            #store an groundtruth of patchclamp_data
            self.patch_clamp_data_ground_truth = self.patch_clamp_data.copy()

            #initialize the patch_clamp_data object with values 
            self.get_baseline_current_info()
            self.allocate_baseline_to_points(initialized=False)
            #self.fill_in_event_events(initialized=False)
            self.add_rolling_mean(initialized=False)
            self.experiment_summary()
            self.event_summary(initialized=False)
            self.find_event_times()
        except (FileNotFoundError, KeyError, IndexError) as e:
            raise ValueError(f"Error loading data from {file_path}: {e}")
    
    def save_figure_dir(self, save_location):
            if save_location is None:
                save_location = os.path.join(os.getcwd(),str(self.molecule_ID), str(self.exp_cond), str(self.exp_ID), 'summary_figures')
                
            else:
                save_location = os.path.join(save_location, str(self.molecule_ID), str(self.exp_cond), str(self.exp_ID), 'summary_figures')

            os.makedirs(save_location, exist_ok=True)
            return save_location

    def get_baseline_current_info(self, baseline_initialization_window=100):
        '''a function to find and store thresholds for the baseline of the experiment'''
        try:
            first_x_current_values = self.patch_clamp_data[self.electrical_current_col].head(baseline_initialization_window)
            self.baseline_current_mean = first_x_current_values.mean()
            self.baseline_current_max = first_x_current_values.max()  ## find the 5th largest if outlier - np.partition(first_x_current_values, -5)[-5]
            self.baseline_current_min = first_x_current_values.min()
            self.noise_std=statistics.stdev(first_x_current_values)
            self.threshold =  self.baseline_current_mean+self.noise_std#max(self.baseline_current_max - self.baseline_current_mean, self.baseline_current_mean - self.baseline_current_min)
            self.size_of_window_for_baseline_current = baseline_initialization_window
            #print(f"threshold={self.threshold}, noise_std={self.noise_std} ,mean={self.baseline_current_mean}")
        except Exception as e:
            raise ValueError(f"Unexpected error in get_baseline_current_info: {e}")
        
    def add_rolling_mean(self, column_used_to_assign_threshold=None, add_rolling_mean_above_threshold_col=False,threshold=None, window_size=20,initialized=True):
        '''a function to assign a rolling mean using convolution for speed as interative loops were to slow'''
        if column_used_to_assign_threshold is None:
            column_used_to_assign_threshold = self.electrical_current_col
        if threshold is None:
            threshold=self.threshold

        data_column = self.patch_clamp_data[column_used_to_assign_threshold].values
        # Drop last windowsize-1 from end eg sliding array of 2 over main array will slide of 1 before the end - only need to be done first pass - as from then
        if not initialized:
            self.patch_clamp_data = self.patch_clamp_data[:-window_size+1]
       
        self.patch_clamp_data['rolling_mean']=np.convolve(data_column , np.ones(window_size), mode='valid')/window_size
        if add_rolling_mean_above_threshold_col:
            self.patch_clamp_data['rolling_mean_above_threshold'] = np.where(self.patch_clamp_data[column_used_to_assign_threshold]>=threshold,1,0)    
                
        # Drop last windowsize-1 from end eg sliding array of 2 over main array will slide of 1 before the end - only need to be done first pass - as from then
        if not initialized:
            self.patch_clamp_data = self.patch_clamp_data[:-window_size+1]
    
    def set_axis_params(self):

        plt.gca().spines['top'].set_visible(False)
        plt.gca().spines['right'].set_visible(False)
        plt.gca().spines['left'].set_linewidth(3)
        plt.gca().spines['bottom'].set_linewidth(3)
        plt.tick_params(axis='both', which='major', width=5, length=10)  
        plt.tick_params(axis='both', which='minor', width=3, length=6)

    def format_yticks_scientific(self):
        # Get the y-ticks from the data
        yticks = plt.gca().get_yticks()

        # Format the y-ticks into scientific notation for yticklabels
        yticklabels = ['{:.1e}'.format(ytick) for ytick in yticks]

        # Set the y-ticks and yticklabels
        plt.yticks(yticks, yticklabels)    
        
    def allocate_baseline_to_points_simple_threshold(self,column_used_to_assign_threshold=None,threshold=None):
        '''a function to assign event 1 or baseline 0 based off one point '''
        if column_used_to_assign_threshold is None:
            column_used_to_assign_threshold = self.electrical_current_col
        if threshold is None:
            threshold=self.threshold
        self.patch_clamp_data['baseline_allocation,0_baseline,1_event_simple_threshold'] = np.where(self.patch_clamp_data[column_used_to_assign_threshold]>=threshold,1,0)    


    def allocate_baseline_to_points(self, column_used_to_assign_threshold=None, threshold=None, window_size=20, percentage_threshold=70,initialized=True):
        '''a function to assign event 1 or baseline 0 based off of a percentage thresehold '''
        if column_used_to_assign_threshold is None:
            column_used_to_assign_threshold = self.electrical_current_col
        if threshold is None:
            threshold = self.threshold

        data_column = self.patch_clamp_data_ground_truth[column_used_to_assign_threshold].values
        # Drop last windowsize-1 from end eg sliding array of 2 over main array will slide of 1 before the end - only need to be done first pass - as from then
        if not initialized:
            self.patch_clamp_data = self.patch_clamp_data[:-window_size+1]
            
        #using a convulsions - sliding an array of one in correct size will give us a look forward window equal to the size of that array
        
        self.patch_clamp_data['sum_of_points_above_baseline_allocation'] = np.convolve(data_column >= threshold , np.ones(window_size), mode='valid').astype(int)
        self.patch_clamp_data['percentage_of_points_above_baseline_allocation'] = ((self.patch_clamp_data['sum_of_points_above_baseline_allocation'] / window_size) * 100).astype(int)
        self.patch_clamp_data['baseline_allocation,0_baseline,1_event'] = ((self.patch_clamp_data['sum_of_points_above_baseline_allocation'] / window_size) >= percentage_threshold / 100).astype(int)
    
    def categorize_data(self,column_title='events_categorised_by_bin_thresholds',labels=None,column_to_catagorize=None, bins=[-float('inf'), 0.5e-10, 5e-10, 4.9e-8, float('inf')],duplicates='drop', ordered=False,bin_range = False):
        '''a function to assign crude bins to the data '''
        if labels == None:
            labels=list(range(0,len(bins)-1))
        if column_to_catagorize is None:
            column_to_catagorize=self.electrical_current_col
        
        self.patch_clamp_data[column_title] = pd.cut(self.patch_clamp_data[column_to_catagorize], bins=bins, labels=labels, duplicates=duplicates, ordered=ordered)
        # Print the bin ranges if applicable
    
        bin_range_dict = {}
        #print(labels)
        for i in range(len(labels)):
            #print('label is  ',labels[i])
            bin_range_dict[labels[i]] = (bins[i], bins[i + 1])
        bin_range_dict = dict(sorted(bin_range_dict.items()))
        self.bin_range_dict = bin_range_dict
        #print(bin_range_dict)


    def fill_in_event_events(self, column_used_to_assign_threshold=None, threshold=None, window_size=20, value_to_fill=1, threshold_to_fill=1,initialized=True):
        '''a function to fill in anmolyless outliers within a event states to try and reduce bouncing between states '''
        
        if column_used_to_assign_threshold is None:
            column_used_to_assign_threshold = 'baseline_allocation,0_baseline,1_event'
        if threshold is None:
            threshold = self.threshold

        data_column = self.patch_clamp_data[column_used_to_assign_threshold].values
        # Drop last windowsize-1 from end eg sliding array of 2 over main array will slide of 1 before the end - only need to be done first pass - as from then
        if not initialized:
            self.patch_clamp_data = self.patch_clamp_data[:-window_size+1]
            
        #using a convulsions - sliding an array of one in correct size will give us a look forward window equal to the size of that array
        
        self.patch_clamp_data[f'number_of_next_{window_size}_in_baseline_event_{value_to_fill}'] = np.convolve(data_column == value_to_fill , np.ones(window_size), mode='valid').astype(int)

        self.patch_clamp_data['baseline_allocation,0_baseline,1_event_filled'] = np.where(self.patch_clamp_data[f'number_of_next_{window_size}_in_baseline_event_{value_to_fill}'] > threshold_to_fill,1,0).astype(int)

    def experiment_summary(self):
        '''initalize experiment summary dataframe'''
        PatchClampData.experiment_summary_dataframe.loc[PatchClampData.count, 'total_experiment_time'] = self.length_of_exp
        PatchClampData.experiment_summary_dataframe.loc[PatchClampData.count, 'molecule_ID'] = self.molecule_ID
        PatchClampData.experiment_summary_dataframe.loc[PatchClampData.count, 'experiment_ID'] = self.exp_ID
        PatchClampData.experiment_summary_dataframe.loc[PatchClampData.count, 'exp_cond'] = self.exp_cond  
        PatchClampData.experiment_summary_dataframe.loc[PatchClampData.count, 'size_of_window_for_baseline_current'] = self.size_of_window_for_baseline_current
        PatchClampData.experiment_summary_dataframe.loc[PatchClampData.count, 'baseline_current_mean'] = self.baseline_current_mean
        PatchClampData.experiment_summary_dataframe.loc[PatchClampData.count, 'baseline_current_max'] = self.baseline_current_max
        PatchClampData.experiment_summary_dataframe.loc[PatchClampData.count, 'baseline_current_min'] = self.baseline_current_min
        PatchClampData.experiment_summary_dataframe.loc[PatchClampData.count, 'baseline_noise_std'] = self.noise_std
        PatchClampData.experiment_summary_dataframe.loc[PatchClampData.count, 'baseline_threshold'] = self.threshold

    def event_summary(self, initialized=True):
        '''initalize event summary dataframe'''
        if not initialized:
            self.event_summary_dataframe = pd.DataFrame()    
              
    def find_event_times(self, column_used_to_assign_if_event='baseline_allocation,0_baseline,1_event',show_groups=False,time_in_event = 0.2,current=None,time=None):
        
        '''this function was not really used in the end but have left logic in incase it is useful
        the idea here is to reduce bouncing between states to make the data more easy to analysis
        but will definitely be need to be optimised'''
        self.event_summary_dataframe = pd.DataFrame()

        self.most_recent_time_in_event_for_event_summary= time_in_event
        if current is None:
            current = self.electrical_current_col
        if time is None:
            time = self.time_col
        column_used_to_update_new_if_event=f"{column_used_to_assign_if_event}_post_filtered"
        self.patch_clamp_data[column_used_to_update_new_if_event]=self.patch_clamp_data[column_used_to_assign_if_event]
        event_group_id = (self.patch_clamp_data[column_used_to_update_new_if_event] != self.patch_clamp_data[column_used_to_update_new_if_event].shift()).cumsum()
        self.patch_clamp_data['event_id'] = event_group_id
        group_data=self.patch_clamp_data.groupby(self.patch_clamp_data['event_id'])
      
        if show_groups:
            print('inital_grouping')
            for group_id, group_df in group_data:
                print(f"Group ID: {group_id}")
                print(len(group_df))
        #0= baseline 1= event
        #first removes 0s assumed to be 1 less than time cut off time  
        prev_event_id=None  
        for group_id, group_df in group_data:   
            if prev_event_id == None:
                prev_event_id = group_id 
            if  group_df[time].max()-group_df[time].min()<= time_in_event and group_df[column_used_to_update_new_if_event].iloc[0] == 0:
                self.patch_clamp_data.loc[group_df.index, 'event_id'] = prev_event_id
            else:
                prev_event_id = group_id
        group_data=self.patch_clamp_data.groupby(self.patch_clamp_data['event_id'])
        if show_groups:
            print('1_grouping')
            for group_id, group_df in group_data:
                print(f"Group ID: {group_id}")
                print(len(group_df))
        #merge contuious 1 events that have been re assigned
        prev_event_id=None
        for group_id, group_df in group_data:
            if prev_event_id == None:
                prev_event_id = group_id
            
            if group_data.get_group(prev_event_id)[column_used_to_update_new_if_event].iloc[0] == 1 and group_df[column_used_to_update_new_if_event].iloc[0] == 1:
                self.patch_clamp_data.loc[group_df.index, 'event_id'] = prev_event_id
                
            else:
                
                prev_event_id = group_id
                
        group_data=self.patch_clamp_data.groupby(self.patch_clamp_data['event_id'])
        #reasign fake 1 events in baseline less than time threshold
        for group_id, group_df in group_data:
            
            if prev_event_id == None:
                prev_event_id = group_id
            if group_df[column_used_to_update_new_if_event].iloc[0] == 1 and group_df[current].mean()<self.baseline_current_max and group_df[time].max()-group_df[time].min()<= time_in_event: 
                self.patch_clamp_data.loc[group_df.index, 'event_id'] = prev_event_id 
            else:
                prev_event_id = group_id
                
        group_data=self.patch_clamp_data.groupby(self.patch_clamp_data['event_id'])
        
        if show_groups:
            print('3_grouping')
            for group_id, group_df in group_data:
                print(f"Group ID: {group_id}")
                print(len(group_df))
        #merge consutive baseline event in to one event
        prev_event_id = None
        for group_id, group_df in group_data:
            if prev_event_id is None:
                prev_event_id = group_id
                #print(f"Initial prev_event_id: {prev_event_id}")
            
            prev_event_value = group_data.get_group(prev_event_id)[column_used_to_update_new_if_event].iloc[0]
            current_event_value = group_df[column_used_to_update_new_if_event].iloc[0]
            
            #print(f"Previous event value: {prev_event_value}, Current event value: {current_event_value}")
            
            if prev_event_value == 0 and current_event_value == 0:
                self.patch_clamp_data.loc[group_df.index, 'event_id'] = prev_event_id
                #print(f"event ID patched with prev_event_id: {prev_event_id}")
            
            if current_event_value == 1:
                prev_event_id = group_id
                #print(f"prev_event_id updated to current group_id: {prev_event_id}")
            
            if prev_event_value == 1 and current_event_value == 0:
                prev_event_id = group_id
                #print(f"prev_event_id updated to current group_id: {prev_event_id}")
            
            if prev_event_value == 0 and current_event_value == 1:
                prev_event_id = group_id
                #print(f"prev_event_id updated to current group_id: {prev_event_id}")
            
            else:
                continue

        group_data=self.patch_clamp_data.groupby(self.patch_clamp_data['event_id'])

        self.event_summary_dataframe['event_start_time'] = group_data[time].first()
        self.event_summary_dataframe['event_end_time'] = group_data[time].last()
        self.event_summary_dataframe['baseline_allocation'] = group_data[column_used_to_update_new_if_event].first()
        self.event_summary_dataframe['time_in_event'] = self.event_summary_dataframe['event_end_time'] - self.event_summary_dataframe['event_start_time']
        self.event_summary_dataframe['mean_current'] = group_data[current].mean()
        
        
        self.event_summary_dataframe.drop_duplicates(keep='first', inplace=True)

    def bin_event_allocation_summary(self, column_used_to_assign_if_event='events_categorised_by_bin_thresholds',current=None,time=None):
        '''a function to store the time spent in each state to output to a csv for use in analysis and figure generation'''
        PatchClampData.experiment_summary_dataframe.loc[PatchClampData.count, f'bin_ranges_and_catagory_allocation'] = str(self.bin_range_dict)
        self.binned_event_summary_dataframe = pd.DataFrame()
        if current is None:
            current='I-mon[A]'
        if time is None:
            time='Time[s](Relative to experiment start)'
        event_group_id = (self.patch_clamp_data[column_used_to_assign_if_event] != self.patch_clamp_data[column_used_to_assign_if_event].shift()).cumsum()
        self.patch_clamp_data['event_id'] = event_group_id
        # Initialize an empty dictionary to store tasks

        group_data=self.patch_clamp_data.groupby(self.patch_clamp_data['event_id'])
        self.binned_event_summary_dataframe['event_start_time'] = group_data[time].first()
        self.binned_event_summary_dataframe['event_end_time'] = group_data[time].last()
        self.binned_event_summary_dataframe['event_allocation'] = group_data[column_used_to_assign_if_event].first()
        self.binned_event_summary_dataframe['time_in_event'] = self.binned_event_summary_dataframe['event_end_time'] - self.binned_event_summary_dataframe['event_start_time']+0.001
        self.binned_event_summary_dataframe['mean_current'] = group_data[current].mean()
        self.binned_event_summary_dataframe.drop_duplicates(keep='first', inplace=True)
        
        self.event_dict = {}
        self.total_time_in_event = {}
        self.time_weighted_mean_current_for_event = {}
        event_name=self.binned_event_summary_dataframe['event_allocation'].astype(int).to_list()
        start_times = group_data[time].first().to_list()
        time_in_events = (self.binned_event_summary_dataframe['event_end_time'] - self.binned_event_summary_dataframe['event_start_time']+0.001).to_list()
        
        # Iterate over the groups formed by groupby
        for event_name, group_data in self.binned_event_summary_dataframe.groupby('event_allocation'):           
            start_times = group_data['event_start_time'].tolist()
            time_in_events = (group_data['event_end_time'] - group_data['event_start_time'] + 0.001).tolist()
            event_durations = [(start, time) for start, time in zip(start_times, time_in_events) if time >= 0.002]
            self.event_dict[event_name] = event_durations
            self.total_time_in_event[event_name] = sum(time_in_events)
            
            # calculate time weighted mean current
            mean_current_values_of_events = group_data['mean_current'].tolist()
            time_weighted_mean_numerator = sum(mean * time for mean, time in zip(mean_current_values_of_events, time_in_events))

            total_time_in_events = sum(time_in_events)
            if total_time_in_events == 0:
                time_weighted_mean =0
                self.time_weighted_mean_current_for_event[event_name] = time_weighted_mean
            else:

                time_weighted_mean = time_weighted_mean_numerator / total_time_in_events
                self.time_weighted_mean_current_for_event[event_name] = time_weighted_mean

                #print("Time-weighted mean:", time_weighted_mean)

            #for each event add the data to the experiment summary dataframe 
            PatchClampData.experiment_summary_dataframe.loc[PatchClampData.count, f'time_weighted_mean_current_for_event_{event_name}'] = "{:.3e}".format(time_weighted_mean)
            PatchClampData.experiment_summary_dataframe.loc[PatchClampData.count, f'time_in_event_{event_name}'] = "{:.3e}".format(sum(time_in_events))
            PatchClampData.experiment_summary_dataframe.loc[PatchClampData.count, f'percentage_of_time_in_event_{event_name}'] = "{:.1f}".format((sum(time_in_events)/self.length_of_exp)*100)    
            if start_times:
                PatchClampData.experiment_summary_dataframe.loc[PatchClampData.count, f'first_time_entering_event_{event_name}'] = start_times[0]
            else:
                PatchClampData.experiment_summary_dataframe.loc[PatchClampData.count, f'first_time_entering_event_{event_name}'] = 0

            if time_in_events:
                PatchClampData.experiment_summary_dataframe.loc[PatchClampData.count, f'first_event_length_{event_name}'] = time_in_events[0]
            else:
                PatchClampData.experiment_summary_dataframe.loc[PatchClampData.count, f'first_event_length_{event_name}'] = 0

    def save_summary_dataframe(self, save_location=None):
        save_location=self.save_figure_dir(save_location)
        self.binned_event_summary_dataframe.to_csv(os.path.join(save_location, 'binned_event_summary.csv'), index=True)
        self.event_summary_dataframe.to_csv(os.path.join(save_location, 'baseline_vs_event_event_summary.csv'), index=True)
#----------------------------------------------------------------------------------
#ploting functionality
#----------------------------------------------------------------------------------   
    def append_image_with_caption_to_doc(self,docx_file, image_path, caption):
        '''a function to output images and captions to supplimnetary info documentation '''
        if not os.path.exists(docx_file):
            # If the document doesn't exist, create a new one
            document = Document()
        else:
            # If the document exists, open it
            document = Document(docx_file)
        
        # Add the image
        document.add_picture(image_path, width=Inches(6))  # Adjust width as needed
        
        # Add the caption
        document.add_paragraph(caption)

        # Save the document
        document.save(docx_file) 
    def basic_static_plot(self, show_fig=True, save_fig=True,save_location=None, current_column_to_plot=None):
        '''a function to plot the data with a meaningful title and without catagorisation'''
        if current_column_to_plot is None:
            current_column_to_plot = self.electrical_current_col
        
        fig, ax= plt.subplots(figsize=(30, 6))
        self.set_axis_params()
        ax.plot(self.patch_clamp_data[self.time_col],
                            self.patch_clamp_data[self.electrical_current_col])
        self.format_yticks_scientific()
        # Plotting of the plot label for the larger figure
        plt.annotate('(a)', xy=(1, 1.05), xycoords='axes fraction', xytext=(-10, 10), textcoords='offset points', ha='right', va='top', bbox=dict(boxstyle="round", alpha=0.1))
        ax.set_title(f"A plot to show current (A) vs time (s) \n for {self.plot_name}")
        ax.set_xlabel('Time relevant to experiment start (s)')
        ax.set_ylabel('Current (A)')
        file_name = f"static_plot.png"
        
        if save_fig:
            save_location=self.save_figure_dir(save_location)
            
            plt.savefig(os.path.join(save_location, file_name), bbox_inches='tight')

        if show_fig:
            plt.show()
        #here
    def generate_interactive_plot(self, save_location=None):
        interactive_fig = figure(width=800, height=400, title=f"Interactive plot to show current (A) vs time (s) \n for molecule id {self.molecule_ID} with experimental conditions {self.exp_cond} experiment id {self.exp_ID}")

        # Create a ColumnDataSource
        data = dict(self.patch_clamp_data)
        source = ColumnDataSource(data=data)
        
        # Plot the data
        interactive_fig.circle(self.time_col, self.electrical_current_col, size=3, color="navy", alpha=0.5, source=source)

        # Create a HoverTool
        hover = HoverTool(tooltips=[
            ("Time", "@{"+str(self.time_col)+"}"),
            ("Current", "@{"+str(self.electrical_current_col)+"}"),
            ("Percentage of points above baseline", "@{" + str('percentage_of_points_above_baseline_allocation') + '}'),
            ('Baseline_allocation', "@{" + str('baseline_allocation,0_baseline,1_event') + '}')
        ])

        # Add hover tool to the plot
        interactive_fig.add_tools(hover)
        # Add y-axis label
        interactive_fig.yaxis.axis_label = "Current (A)"
        
        # Add x-axis label
        interactive_fig.xaxis.axis_label = "Time relevant to experiment start (s)"
        
        file_name = f"interactive_plot_molecule_id_{self.molecule_ID}_exp_cond_{self.exp_cond}_exp_{self.exp_ID}.html"
        
        
        save_location=self.save_figure_dir(save_location)
            
        os.makedirs(save_location, exist_ok=True)
        output_file(os.path.join(save_location, file_name))
        save(interactive_fig)


    
    def static_plot_with_labels(self, current_column_to_plot=None,label_column_title='events_categorised_by_bin_thresholds',show_fig=False, save_fig=True,save_location=None,y_axis_upper_limit=None,y_axis_lower_limit=None):
        '''a function to plot the data with a meaningful title and with catagorisation'''
        
        if current_column_to_plot is None:
            current_column_to_plot = self.electrical_current_col
        fig, ax = plt.subplots(figsize=(30, 6))
        self.set_axis_params()
      


        # Create a custom colormap based on the labels and the color map
     
        scatter = ax.scatter(
            self.patch_clamp_data[self.time_col],
            self.patch_clamp_data[current_column_to_plot],
            c=self.patch_clamp_data[label_column_title], 
            cmap='viridis',
            s=5)
            # Plotting of the plot label for the larger figure
        plt.annotate('(b)', xy=(1, 1.05), xycoords='axes fraction', xytext=(-10, 10), textcoords='offset points', ha='right', va='top', bbox=dict(boxstyle="round", alpha=0.1))
        self.format_yticks_scientific()   
        legend1 = ax.legend(*scatter.legend_elements(), loc='upper left', title="Event")
        ax.add_artist(legend1)
        
        ax.set_xlabel('Time relevant to experiment start (s)')
        ax.set_ylabel('Current (A)')
        ax.set_title(f"A plot to show current (A) vs time (s) \n for {self.plot_name} labeled by event type")
                      
        # if y_axis_upper_limit is not None:
        #     ax.set_ylim(top=y_axis_upper_limit)
            
        # if y_axis_lower_limit is None:
        #     ax.set_ylim(bottom=-self.threshold*10)
        # elif y_axis_lower_limit is not None:
        #     ax.set_ylim(bottom=y_axis_lower_limit)
    
        file_name = f"static_plot_labeled_by_{label_column_title}_y_axis_{current_column_to_plot}.png"
        
        if save_fig:
            save_location=self.save_figure_dir(save_location)
            
            plt.savefig(os.path.join(save_location, file_name), bbox_inches='tight')
            

            plt.close(fig)

            

        if show_fig is True:
            fig.show()            
            

    def static_event_plots(self,append_to_figure_doc=False,existing_docx_file="pc_event_plots.docx",max_current_cutoff= float('inf'), current_column_to_plot=None,label_column_title='events_categorised_by_bin_thresholds',show_fig=True, save_fig=True,save_location=None,y_axis_upper_limit=None,y_axis_lower_limit=None,time_padding=1):
        '''a function to plot all regions of data not in the baseline'''
        if save_fig:
                    if save_location is None:
                        save_location = os.path.join(os.getcwd(),str(self.molecule_ID), str(self.exp_cond), str(self.exp_ID), 'all_events_time_filter')
                        
                    else:
                        save_location = os.path.join(save_location, str(self.molecule_ID), str(self.exp_cond), str(self.exp_ID), 'all_events_time_filter')
                    os.makedirs(save_location, exist_ok=True)
        if current_column_to_plot is None:
            current_column_to_plot = self.electrical_current_col

        # Define colormap
        cmap = ['tab:blue','tab:orange']
        cmap = matplotlib.colors.LinearSegmentedColormap.from_list("", cmap)
 
        for index, row in self.event_summary_dataframe.iterrows():
            if row['baseline_allocation'] == 1:
                event_data = self.patch_clamp_data[(self.patch_clamp_data[self.time_col] >= row['event_start_time']-min(row['time_in_event'],time_padding)) & (self.patch_clamp_data[self.time_col] <= row['event_end_time']+min(row['time_in_event'],time_padding))]
                
                if (event_data[label_column_title] == 0).all():
                    continue

                event_data_to_mask=event_data.copy()
                event_data_mask = event_data_to_mask[self.electrical_current_col] < max_current_cutoff
                index_of_last_point_before_threshold_passed = event_data_mask.idxmin()-1
                end_time=row['event_end_time']+min(row['time_in_event'],time_padding)

                if not event_data_mask.all():
                    # rare instances occurred where the first datapoint is above threshold in this case we display the index of that datapoint and contuine with the loop
                    try:
                        event_data = event_data.loc[:index_of_last_point_before_threshold_passed]
                        end_time = event_data[self.time_col].loc[index_of_last_point_before_threshold_passed]
                    except Exception as e:
                        print(f"An error occurred: {e}")
                        continue 

                fig, ax = plt.subplots(figsize=(30, 6))
                self.set_axis_params()
                scatter=ax.scatter(
                    event_data[self.time_col],
                    event_data[current_column_to_plot],
                    c=event_data[label_column_title], 
                    cmap=cmap,
                    s=3)
                self.format_yticks_scientific()
                legend1 = ax.legend(*scatter.legend_elements(), loc='upper left', title='Event')
                ax.add_artist(legend1)
                
                ax.set_xlabel('Time relevant to experiment start (s)')
                ax.set_ylabel('Current (A)')
                
                ax.set_title(f"A plot to show current (A) vs time (s) for {self.plot_name},\n for the event between {row['event_start_time']:.3f}s and {end_time:.3f}s")
                #plt.xlim(left=row['event_start_time']-min(row['time_in_event'],time_padding), right=row['event_end_time']+min(row['time_in_event'],time_padding))              
                #plt.ylim(min=)    
                file_name = f"from_{row['event_start_time']:.2f}_to_{end_time:.2f}_labeled_by_{label_column_title}_current_filtered_{self.most_recent_time_in_event_for_event_summary}.png"
                image_path=os.path.join(save_location, file_name)
                plt.savefig(image_path, bbox_inches='tight')
                plt.close(fig)
                caption= f"Figure S1 - Graph showing current (A) vs time (s) for {self.plot_name} at +100 mV, for the event between {row['event_start_time']:.3f}s and {end_time:.3f}s."
                with open("event_figure_captions.txt", "a") as caption_file:
                    caption_file.write(f"{caption} \n")
                
                if append_to_figure_doc:
                    self.append_image_with_caption_to_doc(existing_docx_file, image_path, caption)

    def plot_gantt(self, current_column_to_plot=None, save_fig=True,save_location=None):
        '''a function to plot a gantt chart of the differnt states'''
        if current_column_to_plot is None:
            current_column_to_plot = self.electrical_current_col
        # Create a figure and axis
        fig, gnt = plt.subplots(figsize=(30, 6))
        self.set_axis_params()
        # Setting labels for x-axis and y-axis
        gnt.set_xlabel('Time relevant to experiment start (s)')
        gnt.set_ylabel('Event type')
          
        # Setting ticks on y-axis
        gnt.set_yticks([10 * (i + 1) for i in range(len(self.event_dict))])
        # Assuming self.event_dict is your dictionary
        keys = list(self.event_dict.keys())
        print(keys)

        # Assuming gnt is your matplotlib axis
        gnt.set_yticklabels(keys)

        # Plotting of the plot label for the larger figure
        plt.annotate('(c)', xy=(1, 1.05), xycoords='axes fraction', xytext=(-10, 10), textcoords='offset points', ha='right', va='top', bbox=dict(boxstyle="round", alpha=0.1))

        # Define colormap
        cmap = ['tab:blue', 'tab:orange', 'tab:green', 'tab:red', 'tab:purple', 'tab:brown', 'tab:pink', 'tab:gray', 'tab:olive', 'tab:teal']

        gnt.set_title(f"A chart to show the event type vs time (s)\n for {self.plot_name}")

        # Plot tasks from the dictionary
        for i, (event, durations) in enumerate(self.event_dict.items()):
            if not durations:
                continue
            color = cmap[i % len(cmap)]  # Use modulus to loop through colors if tasks exceed colormap length
            if isinstance(durations[0], tuple):  # Multiple bars
                for start, width in durations:
                    gnt.broken_barh([(start, width)], (i * 10 + 5, 9), facecolors=color)
            else:  # Single bar
                gnt.broken_barh([durations], (i * 10 + 5, 9), facecolors=color)

        file_name='gantt_chart.png'

        if save_fig is True:
            save_location=self.save_figure_dir(save_location)
            
            plt.savefig(os.path.join(save_location, file_name), bbox_inches='tight')
                    
            os.makedirs(save_location, exist_ok=True)

        plt.savefig(os.path.join(save_location, file_name), bbox_inches='tight')

        plt.show()     

       
    def time_vs_change_in_current_plot(self,save_fig=True,save_location=None):
        '''a function to plot the rate of change in current'''
        time = self.patch_clamp_data['Time[s](Relative to experiment start)']
        current = self.patch_clamp_data['I-mon[A]']

        # Calculate the change in magnitude
        magnitude_change = current.diff()  
        # Plotting the data
        plt.figure(figsize=(30, 6)) 
        plt.plot(time, magnitude_change)
        # Plotting of the plot label for the larger figure
        plt.annotate('(e)', xy=(1, 1.05), xycoords='axes fraction', xytext=(-10, 10), textcoords='offset points', ha='right', va='top', bbox=dict(boxstyle="round", alpha=0.1))
        plt.xlabel('Time relevant to experiment start (s)')
        plt.ylabel('Change in magnitude of (A)')
        plt.title(f'A plot to show the change in magnitude of current (A) vs time (s)\n for {self.plot_name}')
        self.set_axis_params() 
        self.format_yticks_scientific()       
        if save_fig:
            save_location=self.save_figure_dir(save_location)
            file_name='time_vs_change_in_current_plot.png'
            os.makedirs(save_location, exist_ok=True)
            plt.savefig(os.path.join(save_location, file_name), bbox_inches='tight')
                    
            



    def time_vs_percentage_of_points_above_baseline_allocation_plot(self,save_fig=True,save_location=None):
        '''a function to plot the percentage of points aobve baseline to get an idea of underlying trend'''
        time = self.patch_clamp_data['Time[s](Relative to experiment start)']
        current = self.patch_clamp_data['percentage_of_points_above_baseline_allocation']
        # Plotting the data
        plt.figure(figsize=(30, 6)) 
        plt.plot(time, current)
        # Plotting of the plot label for the larger figure
        plt.annotate('(d)', xy=(1, 1.05), xycoords='axes fraction', xytext=(-10, 10), textcoords='offset points', ha='right', va='top', bbox=dict(boxstyle="round", alpha=0.1))
        plt.xlabel('Time relevant to experiment start (s)')
        plt.ylabel(f'% of points above baseline')
        self.set_axis_params()
        plt.title(f'A plot to show the percentage of points above baseline vs time (s)\n for {self.plot_name}')
        plt.yticks(range(0, 101, 10))
        

        if save_fig:
            save_location=self.save_figure_dir(save_location)
            file_name='percentage_of_points_above_baseline_allocation_vs_time.png'
            os.makedirs(save_location, exist_ok=True)
            plt.savefig(os.path.join(save_location, file_name), bbox_inches='tight')
                    
            
    def create_combined_figure(self,append_to_figure_doc=False,existing_docx_file="pc_exp_summary.docx",offset_increment=20,list_of_figures=['static_plot.png','static_plot_labeled_by_events_categorised_by_bin_thresholds_y_axis_I-mon[A].png','gantt_chart.png','percentage_of_points_above_baseline_allocation_vs_time.png','time_vs_change_in_current_plot.png'],save_location=None):
        ''' a function to join figure together for supplimnetary infomation'''
        
        # Join each filename to the base directory
        filepaths = [os.path.join(self.save_figure_dir(save_location), filename) for filename in list_of_figures]

        # Open all images and store them in a list
        images = [Image.open(filepath) for filepath in filepaths]

        # Calculate total height for the stacked image
        total_height = sum(img.height for img in images)
        
        # Calculate maximum width among the images
        max_width = max(img.width for img in images)
        total_height_and_offset=int(total_height+len(list_of_figures)*(2*total_height/(offset_increment*len(list_of_figures))))
        
        # Create a new blank image with the maximum width and total height
        stacked_image = Image.new('RGB', (max_width,total_height_and_offset ), (255, 255, 255))

        # Create drawing object
        draw = ImageDraw.Draw(stacked_image)

        # Paste each image onto the blank image
        y_offset = 0
        for img in images:

            x_offset = max_width - img.width  # Aligning on the right
            y_offset += int(img.height/20)
            stacked_image.paste(img, (x_offset, y_offset))
            y_offset += img.height + int(img.height/20)
            draw.line([(0, y_offset), (max_width, y_offset)], fill="black", width=10)
        stacked_image_path = os.path.join(self.save_figure_dir(save_location), 'figure_for_paper.png')   
        # Save the stacked image
        stacked_image.save(stacked_image_path)

        # Close all opened images
        for img in images:
            img.close()
        if append_to_figure_doc:
            self.append_image_with_caption_to_doc(existing_docx_file, stacked_image_path, self.stack_image_caption)

    def basic_static_plot_timesearch(self,start_time, end_time, show_fig=True, save_fig=True,figsize=(20,20)):
        '''a function to generate time specfic plots for analysis'''
        event_data = self.patch_clamp_data[((self.patch_clamp_data[self.time_col] >= start_time) & (self.patch_clamp_data[self.time_col] <= end_time))]
        fig, ax = plt.subplots(figsize=figsize)
        self.set_axis_params()
        scatter=ax.scatter(
            event_data[self.time_col],
            event_data[self.electrical_current_col],
            s=3)
        self.format_yticks_scientific()
        
        ax.set_xlabel('Time relevant to experiment start (s)')
        ax.set_ylabel('Current (A)')
        
        ax.set_title(f"A plot to show current (A) vs time (s) for {self.plot_name},\n for the event between {start_time}s and {end_time}s")
        #plt.xlim(left=row['event_start_time']-min(row['time_in_event'],time_padding), right=row['event_end_time']+min(row['time_in_event'],time_padding))              
        #plt.ylim(min=)    
        
        
        if save_fig:            
            directory = 'time_filtered_plots'
        if not os.path.exists(directory):
            os.makedirs(directory)

        file_name = f"from_{start_time}_to_{end_time}_for_{self.molecule_ID}_cond_{self.exp_cond}_exp_{self.exp_ID}.png"
        image_path = os.path.join(directory, file_name)

        plt.savefig(image_path, bbox_inches='tight')

        if show_fig:
            plt.show()


    def create_fig_caption(self):
        bins_in_text_form = ', '.join(f"{key}: {start:.2e} A ≤ Event {key} < {end:.2e} A" for key, (start, end) in self.bin_range_dict.items())
        self.stack_image_caption=f"Figure S1 - Summarising the data from the patch clamp experiment for {self.plot_name} at +100 mV. Events were categorised into the following types: {bins_in_text_form}. For d the percentage of points found above {self.threshold:.1e} (A). The threshold is the mean + standard deviation of the noise of the first 200 datapoints with the percentage calculated in a rolling look ahead window of 20 points."
        with open("stack_figure_captions.txt", "a") as caption_file:
            caption_file.write(f"{self.stack_image_caption} \n")
